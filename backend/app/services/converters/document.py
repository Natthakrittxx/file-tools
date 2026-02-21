import io
from collections.abc import Callable

from docx import Document
from fpdf import FPDF

from app.models import FileFormat


def convert_docx_to_txt(
    input_bytes: bytes,
    source: FileFormat,
    target: FileFormat,
    progress_cb: Callable[[int, str], None] | None = None,
) -> bytes:
    """Extract text from DOCX."""
    if progress_cb:
        progress_cb(50, "Extracting text from document...")
    doc = Document(io.BytesIO(input_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    return text.encode("utf-8")


def convert_txt_to_docx(
    input_bytes: bytes,
    source: FileFormat,
    target: FileFormat,
    progress_cb: Callable[[int, str], None] | None = None,
) -> bytes:
    """Create a DOCX from plain text."""
    if progress_cb:
        progress_cb(50, "Creating document...")
    text = input_bytes.decode("utf-8")
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def convert_txt_to_pdf(
    input_bytes: bytes,
    source: FileFormat,
    target: FileFormat,
    progress_cb: Callable[[int, str], None] | None = None,
) -> bytes:
    """Create a PDF from plain text using fpdf2."""
    if progress_cb:
        progress_cb(50, "Generating PDF...")
    text = input_bytes.decode("utf-8")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    for line in text.split("\n"):
        pdf.multi_cell(0, 6, line)

    return pdf.output()
